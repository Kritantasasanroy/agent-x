"""Render resume / cover-letter text to PDF + DOCX on disk.

reportlab / python-docx are imported lazily. If either is missing (e.g. a minimal
deployment), we fall back to writing a plain-text file with the same base name so the
pipeline keeps working instead of crashing.
"""

from __future__ import annotations

from pathlib import Path

from app.core.config import settings
from app.core.logging import get_logger

log = get_logger("documents")


def _outdir(kind: str) -> Path:
    d = settings.storage_path / "generated" / kind
    d.mkdir(parents=True, exist_ok=True)
    return d


def _txt_fallback(text: str, name: str, kind: str, ext: str) -> str:
    path = _outdir(kind) / f"{name}.{ext}.txt"
    path.write_text(text, encoding="utf-8")
    return str(path)


def to_pdf(text: str, name: str, kind: str = "resumes") -> str:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except Exception:  # noqa: BLE001
        log.warning("reportlab_unavailable_txt_fallback")
        return _txt_fallback(text, name, kind, "pdf")

    path = _outdir(kind) / f"{name}.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=LETTER, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    flow = []
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        first = block.splitlines()[0]
        style = styles["Heading2"] if first.isupper() and len(first) < 60 else styles["BodyText"]
        flow.append(Paragraph(block.replace("\n", "<br/>"), style))
        flow.append(Spacer(1, 8))
    doc.build(flow)
    return str(path)


def to_docx(text: str, name: str, kind: str = "resumes") -> str:
    try:
        from docx import Document
    except Exception:  # noqa: BLE001
        log.warning("python_docx_unavailable_txt_fallback")
        return _txt_fallback(text, name, kind, "docx")

    path = _outdir(kind) / f"{name}.docx"
    document = Document()
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        lines = block.splitlines()
        first = lines[0]
        if first.isupper() and len(first) < 60:
            document.add_heading(first, level=2)
            rest = "\n".join(lines[1:]).strip()
            if rest:
                document.add_paragraph(rest)
        else:
            document.add_paragraph(block)
    document.save(str(path))
    return str(path)


def render_both(text: str, name: str, kind: str = "resumes") -> tuple[str, str]:
    return to_pdf(text, name, kind), to_docx(text, name, kind)
